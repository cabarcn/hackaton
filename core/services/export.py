from pathlib import Path
from django.template.loader import render_to_string
from django.conf import settings
from xhtml2pdf import pisa
from docx import Document
import tempfile

def export_docx(obj):
    d = Document()
    d.add_heading('Reporte de Revisión de Licitación', 0)
    d.add_paragraph(f"Propósito: {obj.proposito}")
    data = obj.observaciones_json or {}
    d.add_heading('Resumen', 1); d.add_paragraph(data.get('resumen_propuesta',''))
    d.add_heading('Observaciones', 1)
    for o in data.get('observaciones', []):
        d.add_paragraph(f"- {o.get('campo')}: {o.get('tipo')} → {o.get('detalle')}")
    d.add_heading('Correcciones sugeridas', 1)
    for c in data.get('correcciones_sugeridas', []):
        d.add_paragraph(f"- {c.get('campo')}: {c.get('sugerencia')}")
    path = Path(settings.BASE_DIR) / f"reporte_postulacion_{obj.id}.docx"
    d.save(path)
    return path

def export_pdf(obj, request):
    html = render_to_string('core/reporte.html', {'obj': obj, 'obs': obj.observaciones_json or {}, 'export': True})
    path = Path(settings.BASE_DIR) / f"reporte_postulacion_{obj.id}.pdf"
    with open(path, "wb") as f:
        pisa.CreatePDF(src=html, dest=f)
    return path

def _render_pdf_from_html(html, base_url):
    # Usa el mismo motor que ya estés usando para PDF (probablemente WeasyPrint).
    try:
        from weasyprint import HTML
        tmp = Path(tempfile.mkstemp(suffix=".pdf")[1])
        HTML(string=html, base_url=base_url).write_pdf(tmp)
        return str(tmp)
    except Exception as e:
        # Fallback: guardar el HTML como .pdf.html para inspección
        tmp = Path(tempfile.mkstemp(suffix=".pdf.html")[1])
        tmp.write_text(html, encoding="utf-8")
        return str(tmp)

def export_renca_pdf(obj, request=None):
    """
    Genera el PDF oficial con campos de SolicitudRenca + tablas (pauta, comisión, normativas, multas)
    """
    ctx = {
        "obj": obj,
        "pauta": obj.pauta_factores.all(),
        "comision": obj.comision_miembros.all(),
        "normativas": obj.normativas.all(),
        "multas": obj.multas.all(),
        "STATIC_URL": settings.STATIC_URL,
    }
    html = render_to_string("core/renca_pdf.html", ctx)
    base_url = request.build_absolute_uri("/") if request else None
    return _render_pdf_from_html(html, base_url)
